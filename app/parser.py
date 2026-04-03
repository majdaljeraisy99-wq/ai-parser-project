from pathlib import Path
import fitz
from docx import Document
from utils import detect_language


def read_txt(file_path: str) -> dict:
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()

    return {
        "text": text,
        "metadata": {
            "file_name": Path(file_path).name,
            "file_type": "txt",
            "pages": 1,
            "language": detect_language(text)
        }
    }


def read_docx(file_path: str) -> dict:
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)

    return {
        "text": text,
        "metadata": {
            "file_name": Path(file_path).name,
            "file_type": "docx",
            "pages": len(doc.paragraphs),
            "language": detect_language(text)
        }
    }


def read_pdf(file_path: str) -> dict:
    doc = fitz.open(file_path)
    pages_text = []

    for page in doc:
        pages_text.append(page.get_text())

    text = "\n".join(pages_text)

    return {
        "text": text,
        "metadata": {
            "file_name": Path(file_path).name,
            "file_type": "pdf",
            "pages": len(doc),
            "language": detect_language(text)
        }
    }


def parse_document(file_path: str) -> dict:
    suffix = Path(file_path).suffix.lower()

    if suffix == ".txt":
        return read_txt(file_path)
    elif suffix == ".docx":
        return read_docx(file_path)
    elif suffix == ".pdf":
        return read_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")